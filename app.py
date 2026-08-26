import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import tempfile
import os
import re

try:
    import pypdf
except ImportError:
    pypdf = None

# --- CONFIGURACIÓN DE LA PÁGINA Y ESTILO CORPORATIVO (AZUL MARINO Y DORADO) ---
st.set_page_config(page_title="U-Formes DOCX", page_icon="🎓", layout="centered")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@1,700&display=swap');

    @keyframes typing {
        from { width: 0 }
        to { width: 100% }
    }

    @keyframes blink {
        50% { border-color: transparent }
    }

    .logo-container {
        display: flex;
        justify-content: center;
        align-items: center;
        margin-top: 20px;
        margin-bottom: 20px;
    }

    .typing-logo {
        font-family: 'Playfair Display', serif;
        font-style: italic;
        font-weight: bold;
        font-size: 50px; 
        background: linear-gradient(45deg, #1B365D, #2C3E50, #D4AF37, #F39C12);
        background-size: 300% 300%;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        animation: gradientShift 5s ease infinite, typing 3s steps(30, end), blink 0.75s step-end infinite;
        white-space: nowrap;
        overflow: hidden;
        border-right: 4px solid #D4AF37;
        width: 10ch;
        text-align: center;
    }

    @keyframes gradientShift {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }

    .stButton>button {
        background-color: #1B365D;
        color: #F39C12;
        font-size: 20px;
        border-radius: 10px;
        padding: 12px 28px;
        width: 100%;
        font-weight: bold;
        border: 2px solid #D4AF37;
        box-shadow: 0px 4px 10px rgba(0,0,0,0.2);
    }
    .stButton>button:hover {
        background-color: #2C3E50;
        color: #FFFFFF;
        border-color: #F39C12;
    }
    </style>
""", unsafe_allow_html=True)

if 'empezar' not in st.session_state:
    st.session_state.empezar = False

if not st.session_state.empezar:
    st.markdown('<div class="logo-container"><div class="typing-logo">U-Formes</div></div>', unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #555; font-size: 1.2rem; font-weight: 500;'>Tu asistente académico inteligente para informes de laboratorio en Word (.docx) editables.</p>", unsafe_allow_html=True)
    st.write("")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("Empezar"):
            st.session_state.empezar = True
            st.rerun()
else:
    st.markdown('<div style="text-align: center;"><span style="font-family: \'Playfair Display\', serif; font-style: italic; font-size: 40px; background: linear-gradient(45deg, #1B365D, #D4AF37); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: bold;">U-Formes</span></div>', unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #D4AF37; font-weight: bold; margin-bottom: 25px;'>Generador Oficial de Informes Académicos - UPS (.DOCX)</p>", unsafe_allow_html=True)

    API_KEY = os.environ.get("GEMINI_API_KEY")
    genai.configure(api_key=API_KEY)
    model = genai.GenerativeModel('gemini-3.6-flash')

    st.markdown("### ⚙️ 1. Estructura del Informe")
    opcion_formato = st.radio("Elige cómo estructurar tu documento:", ["Seleccionar secciones específicas", "Subir mi guía o formato oficial"])

    contenido_guia = ""
    secciones_seleccionadas = []

    if opcion_formato == "Seleccionar secciones específicas":
        opciones_disponibles = [
            "Título", "Objetivo General", "Objetivos Específicos", 
            "Introducción", "Justificación", "Marco Teórico", 
            "Materiales y Equipos", "Metodología", "Desarrollo", 
            "Resultados", "Discusión", "Conclusiones", 
            "Recomendaciones", "Identificación de Riesgos y EPP", 
            "Bibliografía", "Anexos"
        ]
        secciones_seleccionadas = st.multiselect(
            "Selecciona las secciones que deseas incluir:", 
            opciones_disponibles, 
            default=["Objetivo General", "Introducción", "Marco Teórico", "Materiales y Equipos", "Metodología", "Resultados", "Conclusiones", "Bibliografía"]
        )
    else:
        st.info("Sube el archivo de práctica (PDF, Word o TXT). La IA extraerá la estructura base.")
        archivo_guia = st.file_uploader("Sube tu archivo de guía", type=["txt", "pdf", "docx"])
        if archivo_guia is not None:
            extension = archivo_guia.name.split(".")[-1].lower()
            if extension == "txt":
                contenido_guia = archivo_guia.getvalue().decode("utf-8")
            elif extension == "pdf" and pypdf:
                reader = pypdf.PdfReader(archivo_guia)
                for page in reader.pages:
                    contenido_guia += page.extract_text() or ""
            elif extension == "docx":
                doc_temp = Document(archivo_guia)
                for para in doc_temp.paragraphs:
                    contenido_guia += para.text + "\n"
            st.success("¡Guía leída con éxito!")

    st.markdown("### 📋 2. Datos del Estudiante y la Práctica")
    col1, col2 = st.columns(2)
    with col1:
        nombre_estudiante = st.text_input("Nombre del Estudiante", value="")
        laboratorio = st.text_input("Laboratorio / Escenario", value="")
        asignatura = st.text_input("Asignatura", value="")
    with col2:
        nivel_grupo = st.text_input("Nivel / Grupo", value="")
        docente = st.text_input("Nombre del Docente", value="")
        periodo = st.text_input("Periodo Académico", value="")

    col3, col4, col5 = st.columns(3)
    with col3:
        dia = st.number_input("Día", min_value=1, max_value=31, value=26)
    with col4:
        mes = st.number_input("Mes", min_value=1, max_value=12, value=8)
    with col5:
        anio = st.number_input("Año", min_value=2024, max_value=2030, value=2026)
    
    fecha_formateada = f"{int(dia):02d}/{int(mes):02d}/{int(anio)}"
    num_informe = st.text_input("Número de Taller / Práctica", value="")

    tema = st.text_area("Tema del Taller o Práctica", placeholder="Ej: Evaluación de la respuesta glucémica postprandial...", value="")

    st.markdown("### 🖼️ 3. Evidencia Fotográfica (Opcional)")
    st.info("Si no subes imágenes, el apartado de anexos se generará vacío de forma limpia.")
    fotos_anexos = st.file_uploader("Sube imágenes para los anexos", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

    if st.button("Generar Informe Oficial en Word (.docx)"):
        if not tema and not contenido_guia:
            st.warning("Por favor, ingresa al menos el tema o sube una guía de práctica.")
        else:
            with st.spinner("Redactando informe profesional y generando archivo Word editable..."):
                try:
                    if opcion_formato == "Seleccionar secciones específicas":
                        instruccion_est = f"Incluye estrictamente estas secciones: {', '.join(secciones_seleccionadas)}."
                    else:
                        instruccion_est = f"Sigue la estructura de esta guía:\n{contenido_guia}"

                    prompt_maestro = f"""
                    Actúa como un estudiante universitario de excelencia de la Universidad Politécnica Salesiana (UPS).
                    Redacta un informe técnico formal para la asignatura de {asignatura if asignatura else 'la materia'}, sobre el tema: "{tema}".
                    
                    {instruccion_est}
                    
                    INSTRUCCIONES DE FORMATO INNEGOCIABLES Y ESTRICTAS:
                    1. PROHIBIDO poner introducciones duplicadas. Ve directo al grano.
                    2. CERO paréntesis aclaratorios o sobreexplicaciones entre paréntesis dentro del texto. Todo concepto debe integrarse y explicarse directamente en el párrafo.
                    3. No utilices guiones seguidos ni líneas divisorias extrañas hechas con símbolos dentro del texto.
                    4. Tono estrictamente académico, formal, humano y sin muletillas robóticas.
                    """
                    
                    respuesta = model.generate_content(prompt_maestro)
                    informe_texto = respuesta.text
                    informe_texto = re.sub(r'[\$\{\}\\]', '', informe_texto)
                    
                    st.success("¡Informe generado con éxito!")

                    doc = Document()

                    for section in doc.sections:
                        section.top_margin = Inches(0.79)
                        section.bottom_margin = Inches(0.79)
                        section.left_margin = Inches(0.79)
                        section.right_margin = Inches(0.79)

                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Times New Roman'
                    font.size = Pt(11)
                    font.color.rgb = RGBColor(0, 0, 0)
                    style.paragraph_format.line_spacing = 2.0
                    style.paragraph_format.space_after = Pt(6)

                    tabla_encabezado = doc.add_table(rows=6, cols=2)
                    tabla_encabezado.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    celda_uni = tabla_encabezado.cell(0, 0)
                    celda_uni.merge(tabla_encabezado.cell(0, 1))
                    p_uni = celda_uni.paragraphs[0]
                    p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_uni = p_uni.add_run("UNIVERSIDAD POLITÉCNICA SALESIANA")
                    run_uni.bold = True
                    run_uni.font.size = Pt(11)

                    datos_tabla = [
                        (f"Nombre del Estudiante: {nombre_estudiante}", f"Nivel/Grupo: {nivel_grupo}"),
                        (f"Laboratorio/Escenario: {laboratorio}", f"Docente: {docente}"),
                        (f"Asignatura: {asignatura}", f"Periodo Académico: {periodo}"),
                        (f"Fecha: {fecha_formateada}", f"Práctica No.: {num_informe}")
                    ]

                    for r_idx, (col_izq, col_der) in enumerate(datos_tabla, start=1):
                        tabla_encabezado.cell(r_idx, 0).paragraphs[0].add_run(col_izq)
                        tabla_encabezado.cell(r_idx, 1).paragraphs[0].add_run(col_der)

                    celda_tema = tabla_encabezado.cell(5, 0)
                    celda_tema.merge(tabla_encabezado.cell(5, 1))
                    p_tema = celda_tema.paragraphs[0]
                    p_tema.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_tema = p_tema.add_run(f"TEMA DEL TALLER O PRÁCTICA: {tema.upper()}")
                    run_tema.bold = True

                    for row in tabla_encabezado.rows:
                        for cell in row.cells:
                            shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                            cell._tc.get_or_add_tcPr().append(shading_elm)

                    doc.add_paragraph()

                    parrafos = informe_texto.split('\n')
                    for p in parrafos:
                        p = p.strip()
                        if not p:
                            continue
                        
                        texto_limpio = p.replace("*", "").replace("#", "").strip()
                        if "TEMA DEL TALLER" in texto_limpio.upper() or "UNIVERSIDAD POLITÉCNICA SALESIANA" in texto_limpio.upper():
                            continue

                        if "MATERIALES" in texto_limpio.upper() and len(texto_limpio) < 40:
                            h = doc.add_paragraph()
                            run_h = h.add_run(texto_limpio)
                            run_h.bold = True
                            run_h.font.size = Pt(12)
                            
                            t_mat = doc.add_table(rows=2, cols=2)
                            t_mat.alignment = WD_TABLE_ALIGNMENT.CENTER
                            hdr_cells = t_mat.rows[0].cells
                            hdr_cells[0].paragraphs[0].add_run("MATERIALES / EQUIPOS").bold = True
                            hdr_cells[1].paragraphs[0].add_run("USO ESPECÍFICO").bold = True
                            
                            for cell in hdr_cells:
                                shading = parse_xml(r'<w:shd {} w:fill="1B365D"/>'.format(nsdecls('w')))
                                cell._tc.get_or_add_tcPr().append(shading)
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = RGBColor(255, 255, 255)

                            row_cells = t_mat.rows[1].cells
                            row_cells[0].paragraphs[0].add_run("Instrumental de laboratorio y reactivos")
                            row_cells[1].paragraphs[0].add_run("Desarrollo analítico y experimental")
                            doc.add_paragraph()

                        elif "RIESGOS" in texto_limpio.upper() and len(texto_limpio) < 40:
                            h = doc.add_paragraph()
                            run_h = h.add_run(texto_limpio)
                            run_h.bold = True
                            run_h.font.size = Pt(12)
                            
                            t_riesgo = doc.add_table(rows=2, cols=2)
                            t_riesgo.alignment = WD_TABLE_ALIGNMENT.CENTER
                            hdr_cells = t_riesgo.rows[0].cells
                            hdr_cells[0].paragraphs[0].add_run("FACTOR DE RIESGO").bold = True
                            hdr_cells[1].paragraphs[0].add_run("EQUIPO DE PROTECCIÓN (EPP)").bold = True
                            
                            for cell in hdr_cells:
                                shading = parse_xml(r'<w:shd {} w:fill="27AE60"/>'.format(nsdecls('w')))
                                cell._tc.get_or_add_tcPr().append(shading)
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = RGBColor(255, 255, 255)

                            row_cells = t_riesgo.rows[1].cells
                            row_cells[0].paragraphs[0].add_run("Biológico / Químico / Físico")
                            row_cells[1].paragraphs[0].add_run("Mandil, Guantes, Gafas de seguridad")
                            doc.add_paragraph()

                        elif p.startswith("*") or p.startswith("#") or len(texto_limpio) < 50 and any(k in texto_limpio.upper() for k in ["DESCRIPCIÓN", "FUNDAMENTACIÓN", "ACTIVIDADES", "CONCLUSIONES", "RECOMENDACIONES", "BIBLIOGRAFÍA", "ANEXOS", "OBJETIVO", "INTRODUCCIÓN", "JUSTIFICACIÓN", "MARCO", "METODOLOGÍA", "RESULTADOS", "DISCUSIÓN"]):
                            h = doc.add_paragraph()
                            run_h = h.add_run(texto_limpio)
                            run_h.bold = True
                            run_h.font.size = Pt(12)
                        
                        else:
                            p_normal = doc.add_paragraph()
                            p_normal.paragraph_format.left_indent = Inches(0.5)
                            p_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_normal.add_run(texto_limpio)

                    doc.add_heading("EVIDENCIA FOTOGRÁFICA DE LA PRÁCTICA", level=2)
                    if fotos_anexos:
                        for idx, foto in enumerate(fotos_anexos):
                            temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
                            temp_img.write(foto.read())
                            temp_img.close()
                            try:
                                doc.add_picture(temp_img.name, width=Inches(4.5))
                                p_img = doc.add_paragraph(f"Figura {idx+1}. Evidencia de laboratorio.")
                                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception:
                                doc.add_paragraph(f"[Error al cargar la imagen {idx+1}]")
                            os.unlink(temp_img.name)
                    else:
                        p_anexo = doc.add_paragraph("[ Espacio reservado para anexos fotográficos - Sin imágenes adjuntadas ]")
                        p_anexo.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    doc.add_paragraph()

                    tabla_firmas = doc.add_table(rows=2, cols=2)
                    tabla_firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tabla_firmas.cell(0, 0).paragraphs[0].add_run("NOMBRES Y APELLIDOS DEL ESTUDIANTE").bold = True
                    tabla_firmas.cell(0, 1).paragraphs[0].add_run("FIRMA").bold = True
                    tabla_firmas.cell(1, 0).paragraphs[0].add_run(nombre_estudiante)
                    tabla_firmas.cell(1, 1).paragraphs[0].add_run("")

                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                        doc.save(tmp_file.name)
                        tmp_path = tmp_file.name

                    with open(tmp_path, "rb") as docx_file:
                        st.download_button(
                            "📥 Descargar Informe en Word (.docx)", 
                            data=docx_file, 
                            file_name="Informe_UPS_Editable.docx", 
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Ocurrió un error al generar el archivo Word: {e}")
